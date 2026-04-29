import os
import logging
from typing import Optional, List, Dict, Tuple
from pathlib import Path
import psycopg2
from psycopg2.extras import execute_values
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from liftmind.config import settings

# Import embedding utilities for vector search
try:
    from liftmind.embedding_utils import generate_embedding
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False

# Import neural reranker
try:
    from liftmind.reranker import rerank_results as neural_rerank
    RERANKER_AVAILABLE = True
except ImportError:
    RERANKER_AVAILABLE = False

# Import HyDE for vague queries
try:
    from liftmind.hyde import is_vague_symptom_query, should_use_hyde, generate_hypothetical_answer_sync
    HYDE_AVAILABLE = True
except ImportError:
    HYDE_AVAILABLE = False

# Scope qualifier — adds shaft/controller/cabin keyword expansion to retrieval
try:
    from liftmind.scope_qualifier import expand_query_for_scope, apply_scope_penalty, detect_scope
    SCOPE_QUALIFIER_AVAILABLE = True
except ImportError:
    SCOPE_QUALIFIER_AVAILABLE = False

logger = logging.getLogger(__name__)

# Minimum results threshold for cascading fallback
MIN_RESULTS_THRESHOLD = 6


# ============================================================================
# MODEL FAMILIES FOR CROSS-MODEL SEARCH
# ============================================================================

MODEL_FAMILIES = {
    "Elfo": ["Elfo", "Elfo 2", "E3", "Elfo Cabin", "Elfo Electronic",
             "Elfo Hydraulic controller", "Elfo Traction"],
    "Supermec": ["Supermec", "Supermec 2", "Supermec 3"],
    "Freedom": ["Freedom", "Freedom MAXI", "Freedom STEP"],
    "Pollock": ["Pollock (P1)", "Pollock (Q1)"],
}

# Models that support multiple drive types (traction/hydraulic) — need clarification
AMBIGUOUS_DRIVE_MODELS = {
    "Elfo", "Supermec", "Supermec 2", "Supermec 3",
}

# Models that support multiple door types (swing/sliding) — need clarification
AMBIGUOUS_DOOR_MODELS = {
    "Bari", "Elfo", "Supermec", "Supermec 2", "Supermec 3",
}


def get_model_family(model: str) -> Optional[List[str]]:
    """Get the family members for a model."""
    for family_name, members in MODEL_FAMILIES.items():
        if model in members:
            return members
    return None

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    'pdf': ['.pdf'],
    'document': ['.doc', '.docx'],
    'text': ['.txt', '.md'],
    'image': ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp']
}

def get_file_type(filepath: str) -> Optional[str]:
    """Determine file type from extension."""
    ext = Path(filepath).suffix.lower()
    for file_type, extensions in SUPPORTED_EXTENSIONS.items():
        if ext in extensions:
            return file_type
    return None

def get_db_connection():
    """Get database connection."""
    return psycopg2.connect(settings.DATABASE_URL)

def init_database():
    """Initialize the database tables.

    Note: Database schema is now managed by scripts/schema.sql.
    This function just verifies the connection works.
    """
    conn = get_db_connection()
    cur = conn.cursor()

    # Just verify connection works - schema is managed by schema.sql
    cur.execute("SELECT 1")

    cur.close()
    conn.close()
    logger.info("Database connection verified")


def _extract_tables_as_markdown(page, page_num: int) -> List[dict]:
    """
    Extract tables from a pdfplumber page as markdown chunks.

    Keeps tables intact to preserve error code mappings.
    If table > 3000 chars, splits by rows but repeats header.
    """
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed, skipping table extraction")
        return []

    tables = page.extract_tables()
    table_chunks = []

    for table in tables:
        if not table or len(table) < 2:
            continue

        # Build markdown table
        header = table[0] if table else []
        if not any(header):  # Skip empty tables
            continue

        # Create header row
        header_row = "| " + " | ".join(str(h or "").strip() for h in header) + " |"
        separator = "| " + " | ".join("---" for _ in header) + " |"

        # Create data rows
        data_rows = []
        for row in table[1:]:
            row_text = "| " + " | ".join(str(c or "").strip() for c in row) + " |"
            data_rows.append(row_text)

        # Combine into markdown
        full_table = "\n".join([header_row, separator] + data_rows)

        # If table is too large, split by rows (keeping header)
        if len(full_table) > 3000:
            # Split into chunks of ~15 rows each
            for i in range(0, len(data_rows), 15):
                chunk_rows = [header_row, separator] + data_rows[i:i+15]
                table_chunks.append({
                    "content": "\n".join(chunk_rows),
                    "chunk_type": "table",
                    "page_number": page_num,
                })
        else:
            table_chunks.append({
                "content": full_table,
                "chunk_type": "table",
                "page_number": page_num,
            })

    return table_chunks


def ingest_pdf(filepath: str, lift_model: str) -> int:
    """Ingest a PDF into the database with smart table handling."""
    logger.info(f"Ingesting PDF: {filepath} for model {lift_model}")

    # Try pdfplumber first for table-aware extraction
    try:
        import pdfplumber
        return _ingest_pdf_smart(filepath, lift_model)
    except ImportError:
        logger.info("pdfplumber not available, using standard ingestion")
        return _ingest_pdf_standard(filepath, lift_model)


def _ingest_pdf_smart(filepath: str, lift_model: str) -> int:
    """Smart PDF ingestion with table preservation."""
    import pdfplumber

    chunks = []
    filename = os.path.basename(filepath)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    with pdfplumber.open(filepath) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            # Extract tables first (preserve intact)
            table_chunks = _extract_tables_as_markdown(page, page_num)
            for tc in table_chunks:
                tc["lift_model"] = lift_model
                tc["filename"] = filename
                tc["file_type"] = "pdf"
                chunks.append(tc)

            # Extract remaining text (excluding table areas if possible)
            text = page.extract_text() or ""
            if text.strip():
                # Chunk the text content
                text_chunks = splitter.split_text(text)
                for chunk_text in text_chunks:
                    if chunk_text.strip():
                        chunks.append({
                            "lift_model": lift_model,
                            "filename": filename,
                            "file_type": "pdf",
                            "page_number": page_num,
                            "content": chunk_text,
                            "chunk_type": "text"
                        })

    _save_chunks(chunks, filename, lift_model)
    logger.info(f"Smart ingested {filepath}: {len(chunks)} chunks ({sum(1 for c in chunks if c.get('chunk_type') == 'table')} tables)")
    return len(chunks)


def _ingest_pdf_standard(filepath: str, lift_model: str) -> int:
    """Standard PDF ingestion (fallback when pdfplumber unavailable)."""
    loader = PyPDFLoader(filepath)
    pages = loader.load()

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    chunks = []
    for page in pages:
        page_chunks = splitter.split_text(page.page_content)
        for chunk in page_chunks:
            if chunk.strip():  # Skip empty chunks
                chunks.append({
                    "lift_model": lift_model,
                    "filename": os.path.basename(filepath),
                    "file_type": "pdf",
                    "page_number": page.metadata.get("page", 0) + 1,
                    "content": chunk
                })

    _save_chunks(chunks, os.path.basename(filepath), lift_model)
    return len(chunks)


def ingest_docx(filepath: str, lift_model: str) -> int:
    """Ingest a DOCX file into the database."""
    logger.info(f"Ingesting DOCX: {filepath} for model {lift_model}")

    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise

    doc = Document(filepath)

    # Extract all text from paragraphs and tables
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    text = "\n".join(full_text)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    text_chunks = splitter.split_text(text)
    chunks = []
    for i, chunk in enumerate(text_chunks):
        if chunk.strip():
            chunks.append({
                "lift_model": lift_model,
                "filename": os.path.basename(filepath),
                "file_type": "document",
                "page_number": i + 1,
                "content": chunk
            })

    _save_chunks(chunks, os.path.basename(filepath), lift_model)
    return len(chunks)


def ingest_text(filepath: str, lift_model: str) -> int:
    """Ingest a TXT or MD file into the database."""
    logger.info(f"Ingesting text: {filepath} for model {lift_model}")

    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""]
    )

    text_chunks = splitter.split_text(text)
    chunks = []
    for i, chunk in enumerate(text_chunks):
        if chunk.strip():
            chunks.append({
                "lift_model": lift_model,
                "filename": os.path.basename(filepath),
                "file_type": "text",
                "page_number": i + 1,
                "content": chunk
            })

    _save_chunks(chunks, os.path.basename(filepath), lift_model)
    return len(chunks)


def ingest_image(filepath: str, lift_model: str, description: str = "") -> int:
    """
    Index an image file - stores metadata and path for later Claude vision queries.

    Images are NOT OCR'd - instead they're stored so Claude can view them
    when answering questions (useful for wiring diagrams, schematics, etc.)
    """
    logger.info(f"Indexing image: {filepath} for model {lift_model}")

    try:
        from PIL import Image
    except ImportError:
        logger.error("Pillow not installed. Run: pip install Pillow")
        raise

    # Get image dimensions and size
    img = Image.open(filepath)
    width, height = img.size
    img.close()
    file_size = os.path.getsize(filepath)

    # Calculate relative path from manuals dir
    relative_path = os.path.relpath(filepath, settings.MANUALS_DIR)
    filename = os.path.basename(filepath)

    # Auto-generate description from filename if not provided
    if not description:
        # Clean up filename for description
        name_without_ext = Path(filename).stem
        description = name_without_ext.replace('_', ' ').replace('-', ' ')
        description = f"{lift_model} - {description}"

    conn = get_db_connection()
    cur = conn.cursor()

    # Remove existing entry for this image
    cur.execute(
        "DELETE FROM images WHERE filename = %s AND lift_models && %s",
        (filename, [lift_model])
    )

    # Insert new entry (lift_models is an array column)
    cur.execute("""
        INSERT INTO images (lift_models, filename, file_path, relative_path, description, width, height, file_size)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        RETURNING id
    """, ([lift_model], filename, filepath, relative_path, description, width, height, file_size))
    image_id = cur.fetchone()[0]

    conn.commit()
    cur.close()
    conn.close()

    logger.info(f"Indexed image: {filename} ({width}x{height}) id={image_id}")
    return image_id


# Vision prompt for diagram analysis (used by analyze_image_with_vision)
VISION_PROMPT = """Describe this lift/elevator technical diagram for search indexing.

Include:
1. Type of diagram (wiring schematic, mechanical assembly, hydraulic circuit, control panel, installation diagram)
2. All visible terminal numbers, labels, and component names
3. Wire colors if visible
4. Any text, part numbers, or reference codes
5. What system or process this diagram explains

Format as a dense, searchable description like:
"[Type] showing [main subject]. Components: [list]. Terminals: [list]. Labels: [text]. Purpose: [what it explains]."

Be specific and technical. This description will be used for search matching.
Output ONLY the description, no preamble or explanation."""


def analyze_image_with_vision(image_id: int, file_path: str, lift_model: str) -> str:
    """
    Use Claude CLI with vision to generate a searchable description for an image.

    Updates the images table with the generated description.

    Args:
        image_id: Database ID of the image row
        file_path: Path to the image file on disk
        lift_model: Lift model for context

    Returns:
        Generated description, or empty string on failure
    """
    import subprocess

    if not os.path.exists(file_path):
        logger.error(f"Vision analysis: image file not found: {file_path}")
        return ""

    prompt = f"This is a technical diagram for a {lift_model} lift.\n\n{VISION_PROMPT}"

    cmd = ["claude", "-p", prompt, "--image", file_path]

    try:
        logger.info(f"Running vision analysis for image {image_id}: {file_path}")

        process = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120
        )

        if process.returncode != 0:
            logger.error(f"Vision analysis CLI error: {process.stderr}")
            return ""

        description = process.stdout.strip()

        # Clean up any markdown formatting
        if description.startswith("```"):
            lines = description.split("\n")
            description = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])

        if not description:
            logger.warning(f"Vision analysis returned empty description for image {image_id}")
            return ""

        # Update the database
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("UPDATE images SET description = %s WHERE id = %s", (description, image_id))
        conn.commit()
        cur.close()
        conn.close()

        logger.info(f"Vision analysis complete for image {image_id}: {description[:80]}...")
        return description

    except subprocess.TimeoutExpired:
        logger.error(f"Vision analysis timed out for image {image_id}")
        return ""
    except FileNotFoundError:
        logger.error("Claude CLI not found. Make sure 'claude' is in your PATH")
        return ""
    except Exception as e:
        logger.error(f"Vision analysis exception for image {image_id}: {e}")
        return ""


def _save_chunks(chunks: list[dict], filename: str, lift_model: str):
    """Save document chunks to database."""
    conn = get_db_connection()
    cur = conn.cursor()

    # Clear existing data for this file
    cur.execute(
        "DELETE FROM documents WHERE filename = %s AND lift_model = %s",
        (filename, lift_model)
    )

    # Insert new chunks
    for chunk in chunks:
        cur.execute(
            """INSERT INTO documents (lift_model, filename, file_type, page_number, content)
               VALUES (%s, %s, %s, %s, %s)""",
            (chunk["lift_model"], chunk["filename"], chunk["file_type"],
             chunk["page_number"], chunk["content"])
        )

    conn.commit()
    cur.close()
    conn.close()


def ingest_file(filepath: str, lift_model: str) -> int:
    """
    Ingest any supported file type.

    Returns number of chunks/entries created.
    """
    file_type = get_file_type(filepath)

    if file_type == 'pdf':
        return ingest_pdf(filepath, lift_model)
    elif file_type == 'document':
        return ingest_docx(filepath, lift_model)
    elif file_type == 'text':
        return ingest_text(filepath, lift_model)
    elif file_type == 'image':
        return ingest_image(filepath, lift_model)
    else:
        logger.warning(f"Unsupported file type: {filepath}")
        return 0


def update_index_status(lift_model: str, filename: str, file_type: str,
                        relative_path: str, status: str, chunks_count: int = 0,
                        error_message: str = None):
    """Update the indexing status for a file."""
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO index_status (lift_model, filename, file_type, relative_path, status, chunks_count, error_message, indexed_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP)
        ON CONFLICT (lift_model, filename)
        DO UPDATE SET status = %s, chunks_count = %s, error_message = %s, indexed_at = CURRENT_TIMESTAMP
    """, (lift_model, filename, file_type, relative_path, status, chunks_count, error_message,
          status, chunks_count, error_message))

    conn.commit()
    cur.close()
    conn.close()


def get_index_status() -> dict:
    """Get overall indexing status."""
    conn = get_db_connection()
    cur = conn.cursor()

    # Overall stats
    cur.execute("""
        SELECT
            COUNT(*) as total,
            COUNT(*) FILTER (WHERE status = 'success') as success,
            COUNT(*) FILTER (WHERE status = 'error') as errors,
            COUNT(*) FILTER (WHERE status = 'pending') as pending
        FROM index_status
    """)
    row = cur.fetchone()
    overall = {
        "total": row[0],
        "success": row[1],
        "errors": row[2],
        "pending": row[3]
    }

    # By model
    cur.execute("""
        SELECT lift_model,
               COUNT(*) as total,
               COUNT(*) FILTER (WHERE status = 'success') as success,
               SUM(chunks_count) as chunks
        FROM index_status
        GROUP BY lift_model
        ORDER BY lift_model
    """)
    by_model = {}
    for row in cur.fetchall():
        by_model[row[0]] = {
            "total_files": row[1],
            "indexed": row[2],
            "chunks": row[3] or 0
        }

    # Recent errors
    cur.execute("""
        SELECT lift_model, filename, error_message, indexed_at
        FROM index_status
        WHERE status = 'error'
        ORDER BY indexed_at DESC
        LIMIT 10
    """)
    errors = []
    for row in cur.fetchall():
        errors.append({
            "model": row[0],
            "filename": row[1],
            "error": row[2],
            "time": str(row[3]) if row[3] else None
        })

    cur.close()
    conn.close()

    return {
        "overall": overall,
        "by_model": by_model,
        "recent_errors": errors
    }


def _extract_search_terms(query: str) -> str:
    """Extract key technical terms from query and create OR-based tsquery."""
    import re
    # Remove common stop words that aren't useful for technical search
    stop_words = {'is', 'the', 'a', 'an', 'when', 'what', 'how', 'why', 'can', 'could',
                  'would', 'should', 'does', 'do', 'i', 'my', 'me', 'we', 'our', 'help', 'please',
                  'need', 'want', 'have', 'has', 'had', 'be', 'been', 'being', 'am', 'are', 'was',
                  'were', 'it', 'its', 'this', 'that', 'with', 'for', 'on', 'at', 'to', 'from'}

    # Extract words, keeping technical terms (including 'not' for "not levelling" etc.)
    words = re.findall(r'\b[a-zA-Z0-9]+\b', query.lower())
    terms = [w for w in words if w not in stop_words and len(w) > 1]

    if not terms:
        return query  # Fall back to original

    # Create OR-based query (any term matches)
    return ' | '.join(terms)


def _extract_key_terms(query: str) -> List[str]:
    """Extract the most important technical terms from a query for focused search."""
    import re

    # Stop words to filter out
    stop_words = {'is', 'not', 'the', 'a', 'an', 'when', 'what', 'how', 'why', 'can', 'could',
                  'would', 'should', 'does', 'do', 'i', 'my', 'me', 'we', 'our', 'help', 'please',
                  'need', 'want', 'have', 'has', 'had', 'be', 'been', 'being', 'am', 'are', 'was',
                  'were', 'it', 'its', 'this', 'that', 'with', 'for', 'on', 'at', 'to', 'from',
                  'there', 'here', 'just', 'also', 'very', 'too', 'really', 'about', 'after',
                  'before', 'during', 'through', 'into', 'over', 'under', 'again', 'further',
                  'then', 'once', 'only', 'so', 'than', 'but', 'or', 'if', 'because', 'as',
                  'until', 'while', 'of', 'by', 'any', 'some', 'such', 'no', 'nor', 'own',
                  'same', 'few', 'other', 'most', 'more', 'many', 'each', 'every', 'all',
                  'both', 'these', 'those', 'up', 'down', 'out', 'off', 'get', 'got', 'getting'}

    # High-value technical terms (prioritize these)
    technical_terms = {
        'door', 'lock', 'latch', 'interlock', 'motor', 'pump', 'controller', 'board',
        'sensor', 'switch', 'limit', 'safety', 'encoder', 'inverter', 'drive', 'relay',
        'contactor', 'brake', 'valve', 'cylinder', 'piston', 'ram', 'cable', 'rope',
        'leveling', 'levelling', 'floor', 'stop', 'travel', 'speed', 'position',
        'error', 'fault', 'alarm', 'code', 'warning', 'e01', 'e02', 'e03', 'e04',
        'e05', 'e06', 'e07', 'e08', 'e09', 'e10', 'e11', 'e12', 'e13', 'e14', 'e15',
        'w01', 'w02', 'w03', 'w04', 'w05', 'mounting', 'calibrate', 'adjust', 'set',
        'parameter', 'menu', 'display', 'screen', 'navigation', 'setting', 'config',
        'wiring', 'terminal', 'voltage', 'current', 'power', 'supply', 'phase',
        'overload', 'stuck', 'jammed', 'blocked', 'slow', 'fast', 'not', 'closed',
        'open', 'opening', 'closing', 'locking', 'unlocking', 'moving', 'stopping'
    }

    # Extract words
    words = re.findall(r'\b[a-zA-Z0-9]+\b', query.lower())

    # Separate technical and non-technical terms
    tech = [w for w in words if w in technical_terms]
    other = [w for w in words if w not in stop_words and w not in tech and len(w) > 2]

    # Prioritize: technical terms first, then other meaningful words
    return tech + other[:3]  # Max 3 additional non-technical terms


def _generate_search_variants(query: str) -> List[str]:
    """
    Generate multiple search query variants for better coverage.

    Returns a list of query variants:
    1. Key technical terms only (focused search)
    2. Symptom-focused variant (for troubleshooting)
    3. Component-focused variant
    """
    import re

    variants = []
    query_lower = query.lower()

    # Variant 1: Key terms only (2-4 most important words)
    key_terms = _extract_key_terms(query)
    if key_terms and len(key_terms) >= 2:
        # Create OR-based query with key terms
        variants.append(' | '.join(key_terms[:4]))

    # Variant 2: Symptom-focused (for troubleshooting queries)
    symptom_mappings = {
        'not locking': 'lock latch interlock problem',
        'not closing': 'close door gate stuck',
        'not opening': 'open door gate stuck',
        'not moving': 'motor drive stuck travel',
        'not responding': 'controller board stuck frozen',
        'not stopping': 'stop limit overrun brake',
        'not starting': 'start power motor drive',
        'slow': 'speed slow sluggish travel',
        'noise': 'noise sound motor bearing pump',
        'vibration': 'vibration shake motor bearing',
        'stuck': 'stuck jammed blocked frozen',
        'error': 'error fault alarm code',
        'fault': 'fault error alarm code',
    }

    for symptom, expansion in symptom_mappings.items():
        if symptom in query_lower:
            variants.append(expansion)
            break

    # Variant 3: Component + action pairs
    components = ['door', 'lock', 'motor', 'pump', 'controller', 'sensor',
                  'switch', 'brake', 'valve', 'encoder', 'inverter', 'drive']
    actions = ['not', 'stuck', 'error', 'fault', 'slow', 'adjust', 'calibrate',
               'replace', 'check', 'reset', 'configure', 'set']

    found_component = None
    found_action = None

    for comp in components:
        if comp in query_lower:
            found_component = comp
            break

    for act in actions:
        if act in query_lower:
            found_action = act
            break

    if found_component and found_action:
        # Create a focused component+action search
        variants.append(f"{found_component} | {found_action}")

    # Remove duplicates and empty strings
    seen = set()
    unique_variants = []
    for v in variants:
        if v and v not in seen:
            seen.add(v)
            unique_variants.append(v)

    return unique_variants


def _calculate_keyword_overlap(content: str, query: str) -> float:
    """Calculate keyword overlap score between content and query."""
    import re

    # Extract words from both
    content_words = set(re.findall(r'\b[a-zA-Z0-9]+\b', content.lower()))
    query_words = set(re.findall(r'\b[a-zA-Z0-9]+\b', query.lower()))

    # Filter out very short words
    query_words = {w for w in query_words if len(w) > 2}

    if not query_words:
        return 0.0

    # Calculate overlap
    overlap = len(content_words & query_words)
    return min(1.0, overlap / len(query_words))


def _get_source_type_score(result: dict) -> float:
    """Score based on source type (qa_pair > entity > fact > chunk)."""
    # Check content for indicators of high-value sources
    content = result.get('content', '').lower()

    # Procedure indicators (high value)
    procedure_indicators = ['step 1', 'step 2', 'step 3', 'procedure:', 'how to',
                            'first,', 'then,', 'finally,', 'next,', 'warning:',
                            'caution:', 'note:', 'important:']
    if any(ind in content for ind in procedure_indicators):
        return 0.9

    # Error code indicators (high value)
    if any(f'e{i:02d}' in content or f'w{i:02d}' in content for i in range(100)):
        return 0.85

    # Technical specifications (good value)
    spec_indicators = ['voltage', 'current', 'torque', 'pressure', 'speed',
                       'temperature', 'mm', 'nm', 'bar', 'amp', 'volt']
    if any(ind in content for ind in spec_indicators):
        return 0.7

    # Settings/parameters (good value)
    if 'parameter' in content or 'setting' in content or 'menu' in content:
        return 0.65

    # Default chunk score
    return 0.5


def _rerank_results(results: list[dict], query: str) -> list[dict]:
    """
    Re-rank search results for quality.

    Scoring weights:
    - Keyword overlap: 30%
    - Source type (procedure/warning higher): 25%
    - Contains procedure/warning: 20%
    - Original rank score: 25%
    """
    if not results:
        return results

    import re

    # Calculate scores for each result
    scored_results = []
    for r in results:
        content = r.get('content', '')

        # Keyword overlap (30%)
        keyword_score = _calculate_keyword_overlap(content, query) * 0.30

        # Source type score (25%)
        source_score = _get_source_type_score(r) * 0.25

        # Procedure/warning bonus (20%)
        procedure_bonus = 0.0
        content_lower = content.lower()
        if 'warning' in content_lower or 'caution' in content_lower:
            procedure_bonus = 0.20
        elif 'step' in content_lower or 'procedure' in content_lower:
            procedure_bonus = 0.15
        elif 'important' in content_lower or 'note:' in content_lower:
            procedure_bonus = 0.10

        # Original rank (25%) - normalize to 0-1
        original_rank = r.get('rank', 0)
        # PostgreSQL ts_rank typically returns values 0-1, but can be higher
        normalized_rank = min(1.0, original_rank) * 0.25

        # Total score
        total_score = keyword_score + source_score + procedure_bonus + normalized_rank
        r['_rerank_score'] = total_score
        scored_results.append(r)

    # Sort by rerank score (descending)
    scored_results.sort(key=lambda x: x.get('_rerank_score', 0), reverse=True)

    # Remove internal score field
    for r in scored_results:
        if '_rerank_score' in r:
            del r['_rerank_score']

    return scored_results


def _search_facts(cur, query: str, models: List[str] = None, limit: int = 5,
                  exclude_ids: set = None) -> list[dict]:
    """Search facts table with optional model filter.

    Uses OR-based search to find facts matching any key terms.
    Falls back to ILIKE for technical codes (ER8, bbE, H95) when tsquery returns 0.

    Fact confidence scoring:
    - Base score: BM25 rank from ts_rank_cd
    - Multiplied by: confidence_score (default 0.5)
    - Boosted by: feedback rate (times_confirmed_helpful / times_cited * 0.2)

    Formula: final_rank = bm25_rank * confidence_score * (1 + feedback_boost)
    """
    # Guard against short queries that fail with tsquery
    if len(query.strip()) < 2:
        logger.debug("Query too short, skipping facts search")
        return []

    exclude_ids = exclude_ids or set()
    results = []

    # Extract key terms for smarter search
    search_terms = _extract_search_terms(query)

    # Confidence-weighted ranking formula:
    # bm25_rank * confidence_score * (1 + feedback_rate * 0.2)
    # COALESCE handles NULL values with sensible defaults
    confidence_rank_sql = """
        ts_rank_cd(f.content_tsv, to_tsquery('english', %s), 32)
        * COALESCE(f.confidence_score, 0.5)
        * (1.0 + CASE
            WHEN COALESCE(f.times_cited, 0) > 0
            THEN COALESCE(f.times_confirmed_helpful, 0)::float / f.times_cited * 0.2
            ELSE 0.0
          END)
        as rank
    """

    if models:
        cur.execute(f"""
            SELECT f.id, f.lift_models, COALESCE(d.filename, f.source_text, f.section_path) as filename, d.file_type, f.page, f.content,
                   {confidence_rank_sql}
            FROM facts f
            LEFT JOIN documents d ON f.document_id = d.id
            WHERE f.content_tsv @@ to_tsquery('english', %s)
              AND f.lift_models && %s
            ORDER BY rank DESC
            LIMIT %s
        """, (search_terms, search_terms, models, limit))
    else:
        cur.execute(f"""
            SELECT f.id, f.lift_models, COALESCE(d.filename, f.source_text, f.section_path) as filename, d.file_type, f.page, f.content,
                   {confidence_rank_sql}
            FROM facts f
            LEFT JOIN documents d ON f.document_id = d.id
            WHERE f.content_tsv @@ to_tsquery('english', %s)
            ORDER BY rank DESC
            LIMIT %s
        """, (search_terms, search_terms, limit))

    for row in cur.fetchall():
        fact_id = row[0]
        if fact_id not in exclude_ids:
            lift_models = row[1] if row[1] else []
            results.append({
                "id": fact_id,
                "lift_model": lift_models[0] if lift_models else None,
                "filename": row[2],
                "file_type": row[3],
                "page_number": row[4] or 1,
                "content": row[5],
                "rank": row[6],
                "cross_model": False
            })
            exclude_ids.add(fact_id)

    # ILIKE fallback for technical codes.
    # Fires when results are thin (< 3) — not just zero — so that specific
    # part numbers / model codes (SP-200-27, P600M) are always matched even
    # when tsquery finds some generic docs but misses the specific content.
    if len(results) < 3:
        import re as _re
        # Also capture hyphenated codes like SP-200-27 that \b[a-zA-Z0-9]+\b splits
        hyphenated = _re.findall(r'[A-Za-z]{1,5}(?:-\d+)+', query)
        raw_terms = _re.findall(r'\b[a-zA-Z0-9]+\b', query)
        # Look for short alphanumeric codes that tsquery might miss
        code_terms = [t for t in raw_terms if len(t) >= 2 and (
            _re.match(r'^[A-Za-z]+\d+', t) or  # P600M, ER8, H95
            _re.match(r'^\d+[A-Za-z]', t) or    # 95H
            t.upper() != t.lower()               # mixed case codes
        )]
        tech_terms = list(dict.fromkeys(hyphenated + code_terms))  # dedup, preserve order
        if tech_terms:
            existing_content = {r["content"][:80] for r in results}
            ilike_conditions = " OR ".join(["f.content ILIKE %s"] * len(tech_terms))
            ilike_params = [f'%{t}%' for t in tech_terms]

            if models:
                cur.execute(f"""
                    SELECT f.id, f.lift_models, COALESCE(d.filename, f.source_text, f.section_path) as filename, d.file_type, f.page, f.content,
                           0.6 as rank
                    FROM facts f
                    LEFT JOIN documents d ON f.document_id = d.id
                    WHERE ({ilike_conditions})
                      AND f.lift_models && %s
                    LIMIT %s
                """, ilike_params + [models, limit])
            else:
                cur.execute(f"""
                    SELECT f.id, f.lift_models, COALESCE(d.filename, f.source_text, f.section_path) as filename, d.file_type, f.page, f.content,
                           0.6 as rank
                    FROM facts f
                    LEFT JOIN documents d ON f.document_id = d.id
                    WHERE ({ilike_conditions})
                    LIMIT %s
                """, ilike_params + [limit])

            for row in cur.fetchall():
                fact_id = row[0]
                if fact_id not in exclude_ids and row[5][:80] not in existing_content:
                    lift_models = row[1] if row[1] else []
                    results.append({
                        "id": fact_id,
                        "lift_model": lift_models[0] if lift_models else None,
                        "filename": row[2],
                        "file_type": row[3],
                        "page_number": row[4] or 1,
                        "content": row[5],
                        "rank": row[6],
                        "cross_model": False
                    })
                    exclude_ids.add(fact_id)
                    existing_content.add(row[5][:80])

            if results:
                logger.info(f"ILIKE search found {len(results)} facts for terms: {tech_terms}")

    return results


def _search_chunks(cur, query: str, models: List[str] = None, limit: int = 5,
                   exclude_ids: set = None) -> list[dict]:
    """Search chunks table with optional model filter.

    Uses OR-based search to find chunks matching any key terms (consistent with _search_facts).
    Falls back to ILIKE for technical codes when tsquery returns 0.
    """
    # Guard against short queries that fail with tsquery
    if len(query.strip()) < 2:
        logger.debug("Query too short, skipping chunks search")
        return []

    exclude_ids = exclude_ids or set()
    results = []

    # Extract key terms for smarter search (same pattern as _search_facts)
    search_terms = _extract_search_terms(query)

    if models:
        cur.execute("""
            SELECT c.id, c.lift_models, COALESCE(d.filename, c.section_path) as filename, d.file_type, c.page_start, c.content,
                   ts_rank_cd(c.content_tsv, to_tsquery('english', %s), 32) as rank
            FROM chunks c
            LEFT JOIN documents d ON c.document_id = d.id
            WHERE c.content_tsv @@ to_tsquery('english', %s)
              AND c.lift_models && %s
            ORDER BY rank DESC
            LIMIT %s
        """, (search_terms, search_terms, models, limit))
    else:
        cur.execute("""
            SELECT c.id, c.lift_models, COALESCE(d.filename, c.section_path) as filename, d.file_type, c.page_start, c.content,
                   ts_rank_cd(c.content_tsv, to_tsquery('english', %s), 32) as rank
            FROM chunks c
            LEFT JOIN documents d ON c.document_id = d.id
            WHERE c.content_tsv @@ to_tsquery('english', %s)
            ORDER BY rank DESC
            LIMIT %s
        """, (search_terms, search_terms, limit))

    for row in cur.fetchall():
        chunk_id = f"chunk_{row[0]}"
        if chunk_id not in exclude_ids:
            lift_models = row[1] if row[1] else []
            results.append({
                "id": chunk_id,
                "lift_model": lift_models[0] if lift_models else None,
                "filename": row[2],
                "file_type": row[3],
                "page_number": row[4] or 1,
                "content": row[5],
                "rank": row[6],
                "cross_model": False
            })
            exclude_ids.add(chunk_id)

    # ILIKE fallback for technical codes.
    # Fires when results are thin (< 3) — not just zero — so that specific
    # part numbers / model codes (SP-200-27, P600M) are always matched even
    # when tsquery finds some generic docs but misses the specific content.
    if len(results) < 3:
        import re as _re
        # Also capture hyphenated codes like SP-200-27 that \b[a-zA-Z0-9]+\b splits
        hyphenated = _re.findall(r'[A-Za-z]{1,5}(?:-\d+)+', query)
        raw_terms = _re.findall(r'\b[a-zA-Z0-9]+\b', query)
        code_terms = [t for t in raw_terms if len(t) >= 2 and (
            _re.match(r'^[A-Za-z]+\d+', t) or   # P600M, ACS850, E07
            _re.match(r'^\d+[A-Za-z]', t)        # 24V, 95H
        )]
        tech_terms = list(dict.fromkeys(hyphenated + code_terms))  # dedup, preserve order
        if tech_terms:
            existing_content = {r["content"][:80] for r in results}
            ilike_conditions = " OR ".join(["c.content ILIKE %s"] * len(tech_terms))
            ilike_params = [f'%{t}%' for t in tech_terms]

            if models:
                cur.execute(f"""
                    SELECT c.id, c.lift_models, COALESCE(d.filename, c.section_path) as filename, d.file_type, c.page_start, c.content,
                           0.6 as rank
                    FROM chunks c
                    LEFT JOIN documents d ON c.document_id = d.id
                    WHERE ({ilike_conditions})
                      AND c.lift_models && %s
                    LIMIT %s
                """, ilike_params + [models, limit])
            else:
                cur.execute(f"""
                    SELECT c.id, c.lift_models, COALESCE(d.filename, c.section_path) as filename, d.file_type, c.page_start, c.content,
                           0.6 as rank
                    FROM chunks c
                    LEFT JOIN documents d ON c.document_id = d.id
                    WHERE ({ilike_conditions})
                    LIMIT %s
                """, ilike_params + [limit])

            for row in cur.fetchall():
                chunk_id = f"chunk_{row[0]}"
                if chunk_id not in exclude_ids and row[5][:80] not in existing_content:
                    lift_models = row[1] if row[1] else []
                    results.append({
                        "id": chunk_id,
                        "lift_model": lift_models[0] if lift_models else None,
                        "filename": row[2],
                        "file_type": row[3],
                        "page_number": row[4] or 1,
                        "content": row[5],
                        "rank": row[6],
                        "cross_model": False
                    })
                    exclude_ids.add(chunk_id)
                    existing_content.add(row[5][:80])

            if results:
                logger.info(f"ILIKE search found {len(results)} chunks for terms: {tech_terms}")

    return results


def _run_single_search(cur, query: str, models: List[str] = None, limit: int = 5,
                        seen_ids: set = None) -> list[dict]:
    """Run a single search iteration against facts and chunks."""
    seen_ids = seen_ids or set()
    results = []

    # Search facts first (higher priority)
    results.extend(_search_facts(cur, query, models, limit, seen_ids))

    # Then chunks if needed
    if len(results) < limit:
        remaining = limit - len(results)
        results.extend(_search_chunks(cur, query, models, remaining, seen_ids))

    return results


# ============================================================================
# HYBRID SEARCH: BM25 + VECTOR WITH RRF FUSION
# ============================================================================

def _search_vector_facts(cur, embedding: List[float], models: List[str] = None,
                         limit: int = 30) -> List[dict]:
    """
    Search facts table using pgvector cosine similarity.

    Args:
        cur: Database cursor
        embedding: 384-dimensional embedding vector
        models: Optional list of lift models to filter
        limit: Max results to return

    Returns:
        List of dicts with id, content, similarity score, metadata
    """
    if not embedding:
        return []

    # Format embedding for pgvector
    embedding_str = '[' + ','.join(str(x) for x in embedding) + ']'

    results = []

    if models:
        cur.execute("""
            SELECT f.id, f.lift_models, COALESCE(d.filename, f.source_text, f.section_path) as filename, d.file_type, f.page, f.content,
                   1 - (f.embedding <=> %s::vector) as similarity
            FROM facts f
            LEFT JOIN documents d ON f.document_id = d.id
            WHERE f.embedding IS NOT NULL
              AND f.lift_models && %s
            ORDER BY f.embedding <=> %s::vector
            LIMIT %s
        """, (embedding_str, models, embedding_str, limit))
    else:
        cur.execute("""
            SELECT f.id, f.lift_models, COALESCE(d.filename, f.source_text, f.section_path) as filename, d.file_type, f.page, f.content,
                   1 - (f.embedding <=> %s::vector) as similarity
            FROM facts f
            LEFT JOIN documents d ON f.document_id = d.id
            WHERE f.embedding IS NOT NULL
            ORDER BY f.embedding <=> %s::vector
            LIMIT %s
        """, (embedding_str, embedding_str, limit))

    for row in cur.fetchall():
        lift_models = row[1] if row[1] else []
        results.append({
            "id": row[0],
            "lift_model": lift_models[0] if lift_models else None,
            "filename": row[2],
            "file_type": row[3],
            "page_number": row[4] or 1,
            "content": row[5],
            "similarity": row[6],
            "source": "vector_facts",
            "cross_model": False
        })

    return results


def _search_vector_chunks(cur, embedding: List[float], models: List[str] = None,
                          limit: int = 30) -> List[dict]:
    """
    Search chunks table using pgvector cosine similarity.
    """
    if not embedding:
        return []

    embedding_str = '[' + ','.join(str(x) for x in embedding) + ']'

    results = []

    if models:
        cur.execute("""
            SELECT c.id, c.lift_models, COALESCE(d.filename, c.section_path) as filename, d.file_type, c.page_start, c.content,
                   1 - (c.embedding <=> %s::vector) as similarity
            FROM chunks c
            LEFT JOIN documents d ON c.document_id = d.id
            WHERE c.embedding IS NOT NULL
              AND c.lift_models && %s
            ORDER BY c.embedding <=> %s::vector
            LIMIT %s
        """, (embedding_str, models, embedding_str, limit))
    else:
        cur.execute("""
            SELECT c.id, c.lift_models, COALESCE(d.filename, c.section_path) as filename, d.file_type, c.page_start, c.content,
                   1 - (c.embedding <=> %s::vector) as similarity
            FROM chunks c
            LEFT JOIN documents d ON c.document_id = d.id
            WHERE c.embedding IS NOT NULL
            ORDER BY c.embedding <=> %s::vector
            LIMIT %s
        """, (embedding_str, embedding_str, limit))

    for row in cur.fetchall():
        lift_models = row[1] if row[1] else []
        results.append({
            "id": f"chunk_{row[0]}",
            "lift_model": lift_models[0] if lift_models else None,
            "filename": row[2],
            "file_type": row[3],
            "page_number": row[4] or 1,
            "content": row[5],
            "similarity": row[6],
            "source": "vector_chunks",
            "cross_model": False
        })

    return results


def _search_vector(embedding: List[float], models: List[str] = None,
                   limit: int = 30) -> List[dict]:
    """
    Combined vector search across facts and chunks.

    Args:
        embedding: 384-dimensional embedding vector from semantic query
        models: Optional list of lift models to filter
        limit: Max results to return

    Returns:
        List of results sorted by similarity (highest first)
    """
    conn = get_db_connection()
    cur = conn.cursor()

    try:
        # Search both facts and chunks
        fact_results = _search_vector_facts(cur, embedding, models, limit)
        chunk_results = _search_vector_chunks(cur, embedding, models, limit)

        # Combine and sort by similarity
        all_results = fact_results + chunk_results
        all_results.sort(key=lambda x: x.get('similarity', 0), reverse=True)

        return all_results[:limit]

    finally:
        cur.close()
        conn.close()


def _search_bm25(keyword_queries: List[str], models: List[str] = None,
                 limit: int = 30) -> List[dict]:
    """
    BM25/Full-text search across facts and chunks.

    Uses OR-based tsvector search + exact phrase ILIKE for codes/numbers.

    Args:
        keyword_queries: List of keyword query strings
        models: Optional list of lift models to filter
        limit: Max results to return

    Returns:
        List of results with BM25 rank scores
    """
    conn = get_db_connection()
    cur = conn.cursor()

    try:
        all_results = []
        seen_ids = set()

        # Phase 1: Exact phrase ILIKE search (high priority)
        # This catches numeric codes, PLC addresses, and specific labels
        # that PostgreSQL tsvector strips (numbers, punctuation)
        for query in keyword_queries[:5]:
            if len(all_results) >= limit:
                break
            # Only use ILIKE for short, specific phrases (likely codes/labels)
            if len(query) <= 30:
                ilike_results = _search_chunks_ilike(cur, query, models, 10, seen_ids)
                for r in ilike_results:
                    r['source'] = 'bm25_ilike'
                    r['rank'] = r.get('rank', 0) + 1.0  # Boost exact matches
                all_results.extend(ilike_results)

        # Phase 2: Standard tsvector search for each keyword query variant
        for query in keyword_queries[:5]:  # Increased from 3 to 5
            if len(all_results) >= limit:
                break

            remaining = limit - len(all_results)

            # Search facts
            fact_results = _search_facts(cur, query, models, remaining, seen_ids)
            for r in fact_results:
                r['source'] = 'bm25_facts'
            all_results.extend(fact_results)

            # Search chunks
            if len(all_results) < limit:
                remaining = limit - len(all_results)
                chunk_results = _search_chunks(cur, query, models, remaining, seen_ids)
                for r in chunk_results:
                    r['source'] = 'bm25_chunks'
                all_results.extend(chunk_results)

        return all_results[:limit]

    finally:
        cur.close()
        conn.close()


def _search_chunks_ilike(cur, query: str, models: List[str] = None, limit: int = 10,
                         exclude_ids: set = None) -> list[dict]:
    """Exact phrase ILIKE search on chunks for codes/numbers tsvector misses."""
    exclude_ids = exclude_ids or set()
    results = []

    # Clean the query for ILIKE (escape % and _)
    clean_query = query.replace('%', '\\%').replace('_', '\\_')
    ilike_pattern = f'%{clean_query}%'

    if models:
        cur.execute("""
            SELECT c.id, c.lift_models, d.filename, d.file_type, c.page_start, c.content
            FROM chunks c
            LEFT JOIN documents d ON c.document_id = d.id
            WHERE c.content ILIKE %s
              AND c.lift_models && %s
            LIMIT %s
        """, (ilike_pattern, models, limit))
    else:
        cur.execute("""
            SELECT c.id, c.lift_models, d.filename, d.file_type, c.page_start, c.content
            FROM chunks c
            LEFT JOIN documents d ON c.document_id = d.id
            WHERE c.content ILIKE %s
            LIMIT %s
        """, (ilike_pattern, limit))

    for row in cur.fetchall():
        chunk_id = row[0]
        if chunk_id not in exclude_ids:
            lift_models = row[1] if row[1] else []
            results.append({
                "id": chunk_id,
                "lift_model": lift_models[0] if lift_models else None,
                "filename": row[2],
                "file_type": row[3],
                "page_number": row[4] or 1,
                "content": row[5],
                "rank": 0.5,  # Base rank for ILIKE results
                "cross_model": False
            })
            exclude_ids.add(chunk_id)

    return results


def reciprocal_rank_fusion(bm25_results: List[dict], vector_results: List[dict],
                           k: int = None, query_type: str = None) -> List[dict]:
    """
    Merge BM25 and Vector search results using Reciprocal Rank Fusion.

    RRF formula: score = sum(1 / (k + rank)) for each result list

    Args:
        bm25_results: Results from BM25/keyword search
        vector_results: Results from vector/semantic search
        k: RRF constant (overrides query_type-based k if specified)
        query_type: One of "fault_code", "specification", "procedure", "general"
                   Used to select optimal k value (lower k = favor BM25 for exact matches)

    Returns:
        Merged and sorted list of results with RRF scores
    """
    # Use configurable k: explicit k > query_type k > default
    if k is None:
        if query_type and query_type in settings.RRF_K_BY_QUERY_TYPE:
            k = settings.RRF_K_BY_QUERY_TYPE[query_type]
        else:
            k = settings.RRF_K_PARAMETER
    scores: Dict[str, float] = {}
    items: Dict[str, dict] = {}

    # Score BM25 results
    for rank, item in enumerate(bm25_results):
        item_id = str(item.get('id', ''))
        if not item_id:
            continue
        scores[item_id] = scores.get(item_id, 0) + 1 / (k + rank + 1)
        items[item_id] = item

    # Score Vector results
    for rank, item in enumerate(vector_results):
        item_id = str(item.get('id', ''))
        if not item_id:
            continue
        scores[item_id] = scores.get(item_id, 0) + 1 / (k + rank + 1)
        # Keep vector item if not already seen (preserves more metadata)
        if item_id not in items:
            items[item_id] = item

    # Sort by fused score
    sorted_ids = sorted(scores.keys(), key=lambda x: scores[x], reverse=True)

    # Build result list with RRF scores
    results = []
    for item_id in sorted_ids:
        if item_id in items:
            result = items[item_id].copy()
            result['rrf_score'] = scores[item_id]
            results.append(result)

    return results


def hybrid_search(keyword_queries: List[str], semantic_query: str,
                  models: List[str] = None, limit: int = 30,
                  use_hyde: bool = True, query_type: str = None) -> List[dict]:
    """
    Run BM25 + Vector search in parallel, merge with RRF.

    This is the core hybrid search function that combines:
    1. BM25 keyword search (exact term matching)
    2. Vector semantic search (meaning-based)
    3. Reciprocal Rank Fusion (RRF) for merging
    4. HyDE enhancement for vague symptom queries

    Args:
        keyword_queries: List of keyword queries for BM25
        semantic_query: Semantic query for vector search
        models: Optional list of lift models to filter
        limit: Max results to return
        use_hyde: Whether to use HyDE for vague queries
        query_type: One of "fault_code", "specification", "procedure", "general"
                   Used to select optimal RRF k value

    Returns:
        RRF-fused results combining both search strategies
    """
    # Check if embeddings are available
    if not EMBEDDINGS_AVAILABLE:
        logger.warning("Embeddings not available, falling back to BM25-only search")
        return _search_bm25(keyword_queries, models, limit)

    # HyDE enhancement for vague symptom queries
    search_text = semantic_query
    if use_hyde and HYDE_AVAILABLE and should_use_hyde(semantic_query):
        model_hint = models[0] if models else None
        hypothetical = generate_hypothetical_answer_sync(semantic_query, model_hint)
        if hypothetical:
            # Combine original query with hypothetical for richer embedding
            search_text = f"{semantic_query} {hypothetical}"
            logger.debug(f"HyDE activated, combined search text: {search_text[:200]}...")

    # Generate embedding for semantic query (possibly HyDE-enhanced)
    embedding = generate_embedding(search_text)

    if embedding is None:
        logger.warning("Failed to generate embedding, falling back to BM25-only search")
        return _search_bm25(keyword_queries, models, limit)

    # Run both searches
    logger.info(f"Running hybrid search: BM25 queries={keyword_queries[:2]}, models={models}, query_type={query_type}")

    bm25_results = _search_bm25(keyword_queries, models, limit=100)
    vector_results = _search_vector(embedding, models, limit=100)

    logger.info(f"BM25 returned {len(bm25_results)}, Vector returned {len(vector_results)}")

    # For technical queries (fault_code, specification), boost BM25 by adding top results twice
    # This gives exact keyword matches more weight in RRF fusion
    if query_type in ("fault_code", "specification") and bm25_results:
        # Add top 50% of BM25 results again to increase their RRF weight
        boost_count = len(bm25_results) // 2
        if boost_count > 0:
            bm25_results = bm25_results + bm25_results[:boost_count]
            logger.info(f"BM25 boosted for {query_type}: added {boost_count} duplicate entries")

    # Fuse results with RRF (k selected based on query_type for optimal BM25/vector balance)
    fused = reciprocal_rank_fusion(bm25_results, vector_results, query_type=query_type)

    # Get actual k used for logging
    actual_k = settings.RRF_K_BY_QUERY_TYPE.get(query_type, settings.RRF_K_PARAMETER) if query_type else settings.RRF_K_PARAMETER
    logger.info(f"RRF fusion produced {len(fused)} unique results (k={actual_k}, query_type={query_type})")

    return fused[:limit]


def search_with_filters(keyword_queries: List[str], semantic_query: str,
                        filters: dict, limit: int = 30, query_type: str = None) -> List[dict]:
    """
    Apply metadata filters BEFORE search with cascading fallback.

    Cascading fallback logic:
    1. Strict model filter - if specified
    2. Model family filter - if < MIN_RESULTS_THRESHOLD results
    3. All models - if still < MIN_RESULTS_THRESHOLD results

    This prevents zero-result failures when the interceptor guesses wrong.

    Args:
        keyword_queries: List of keyword queries for BM25
        semantic_query: Semantic query for vector search
        filters: Dict with model, component, error_code from interceptor
        limit: Max results to return
        query_type: One of "fault_code", "specification", "procedure", "general"
                   Used to select optimal RRF k value

    Returns:
        Search results with cross_model flag set for fallback results
    """
    models = filters.get('model')  # Already a list from interceptor
    error_code = filters.get('error_code')
    component = filters.get('component')

    # Inject component as additional BM25 keyword if not already present
    if component and component.lower() not in ' '.join(keyword_queries).lower():
        keyword_queries = keyword_queries + [component.lower()]

    # Scope-qualifier expansion (e.g. "in the shaft" adds shaft-related keywords)
    if SCOPE_QUALIFIER_AVAILABLE:
        scope = detect_scope(semantic_query)
        if scope:
            keyword_queries = expand_query_for_scope(semantic_query, keyword_queries)
            logger.info(f"Scope qualifier '{scope['scope']}' detected — added keywords")

    # Inject third-party equipment into search terms
    third_party = filters.get('third_party_equipment')
    if third_party:
        if third_party.lower() not in ' '.join(keyword_queries).lower():
            keyword_queries = keyword_queries + [third_party.lower()]
        semantic_query = f"{semantic_query} {third_party}"

    # Fast path: direct error code lookup
    if error_code:
        entity_results = _lookup_error_code(error_code, models)
        if entity_results:
            logger.info(f"Fast path: found {len(entity_results)} results for error code {error_code}")
            return entity_results

    # Step 1: Try STRICT filtered search
    results = hybrid_search(keyword_queries, semantic_query, models, limit, query_type=query_type)

    # Step 2: CASCADING FALLBACK
    if len(results) < MIN_RESULTS_THRESHOLD and models is not None:
        logger.warning(f"Strict filter {models} returned {len(results)} results. Widening search...")

        # Get model family first
        if models and len(models) > 0:
            family = get_model_family(models[0])
            if family:
                # Search other family members
                other_family = [m for m in family if m not in models]
                if other_family:
                    family_results = hybrid_search(keyword_queries, semantic_query,
                                                   other_family, limit - len(results),
                                                   query_type=query_type)
                    for r in family_results:
                        r['cross_model'] = True
                    results.extend(family_results)

                    if len(results) >= MIN_RESULTS_THRESHOLD:
                        logger.info(f"Family search brought results to {len(results)}")
                        return results[:limit]

        # Last resort: search ALL models
        if len(results) < MIN_RESULTS_THRESHOLD:
            logger.warning("Family search insufficient, searching all models")
            broad_results = hybrid_search(keyword_queries, semantic_query,
                                          models=None, limit=limit - len(results),
                                          query_type=query_type)
            for r in broad_results:
                r['cross_model'] = True
            results.extend(broad_results)

    return results[:limit]


def _lookup_error_code(error_code: str, models: List[str] = None) -> Optional[dict]:
    """
    Fast lookup for error codes in entities/facts tables.

    Returns single result if exact error code match found.
    """
    conn = get_db_connection()
    cur = conn.cursor()

    try:
        # Normalize error code (E23 -> E23, e23 -> E23)
        error_code_upper = error_code.upper()

        # Search in facts for error code content
        if models:
            cur.execute("""
                SELECT f.id, f.lift_models, d.filename, d.file_type, f.page, f.content
                FROM facts f
                LEFT JOIN documents d ON f.document_id = d.id
                WHERE UPPER(f.content) LIKE %s
                  AND f.lift_models && %s
                ORDER BY f.id
                LIMIT 5
            """, (f'%{error_code_upper}%', models))
        else:
            cur.execute("""
                SELECT f.id, f.lift_models, d.filename, d.file_type, f.page, f.content
                FROM facts f
                LEFT JOIN documents d ON f.document_id = d.id
                WHERE UPPER(f.content) LIKE %s
                ORDER BY f.id
                LIMIT 5
            """, (f'%{error_code_upper}%',))

        rows = cur.fetchall()
        if rows:
            # Return all matching results for richer context
            results = []
            for row in rows:
                lift_models = row[1] if row[1] else []
                results.append({
                    "id": row[0],
                    "lift_model": lift_models[0] if lift_models else None,
                    "filename": row[2],
                    "file_type": row[3],
                    "page_number": row[4] or 1,
                    "content": row[5],
                    "source": "error_code_lookup",
                    "cross_model": False
                })
            return results

        return None

    finally:
        cur.close()
        conn.close()


def search_documents_hybrid(interceptor_result: dict, limit: int = 5, query_type: str = None) -> list[dict]:
    """
    Search documents using hybrid search with interceptor results.

    This is the primary search function for the new RAG pipeline.
    Uses BM25 + Vector search with RRF fusion and cascading fallback.

    Args:
        interceptor_result: Dict from slang_interceptor with:
            - filters: {model: [...], component: str, error_code: str}
            - keyword_queries: List of keyword queries
            - semantic_query: Semantic query string
            - deep_dive: bool
            - query_type: Optional query type for RRF k selection
        limit: Max results to return
        query_type: Override query_type (defaults to interceptor_result['query_type'])

    Returns:
        List of search results with metadata
    """
    filters = interceptor_result.get('filters', {})
    keyword_queries = interceptor_result.get('keyword_queries', [])
    semantic_query = interceptor_result.get('semantic_query', '')

    # Get query_type from interceptor if not explicitly provided
    if query_type is None:
        query_type = interceptor_result.get('query_type')

    # Ensure we have something to search with
    if not keyword_queries and not semantic_query:
        logger.warning("No keyword_queries or semantic_query provided")
        return []

    # If only semantic query, use it for both
    if not keyword_queries:
        keyword_queries = [semantic_query]

    # If only keywords, join them for semantic
    if not semantic_query and keyword_queries:
        semantic_query = ' '.join(keyword_queries)

    # Run hybrid search with cascading fallback (pass query_type for dynamic RRF k)
    results = search_with_filters(keyword_queries, semantic_query, filters, limit * 6, query_type=query_type)

    # Re-rank for quality - use neural reranker if available
    if RERANKER_AVAILABLE and len(results) > 0:
        logger.info(f"Neural reranking {len(results)} results...")
        results = neural_rerank(semantic_query, results, top_k=limit * 2)
    else:
        # Fall back to basic keyword-based reranking
        results = _rerank_results(results, semantic_query)

    # --- Document-level fact expansion ---
    # When chunks from a document are found, also pull in related facts
    # from the same document. This catches short, specific facts (e.g.
    # "screw 2 = latching action") that BM25 ranks too low on their own.
    try:
        doc_filenames = set()
        seen_fact_ids = set()
        for r in results[:limit]:
            fn = r.get('filename')
            if fn:
                doc_filenames.add(fn)
            fact_id = r.get('id')
            if fact_id:
                seen_fact_ids.add(fact_id)

        if doc_filenames:
            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute("""
                    SELECT f.id, f.content, d.filename, f.page, f.lift_models
                    FROM facts f
                    LEFT JOIN documents d ON f.document_id = d.id
                    WHERE d.filename = ANY(%s)
                    LIMIT 20
                """, (list(doc_filenames),))

                expansion_count = 0
                for row in cur.fetchall():
                    fact_id, content, filename, page, models = row
                    if fact_id not in seen_fact_ids:
                        results.append({
                            'id': fact_id,
                            'content': content,
                            'filename': filename,
                            'page_number': page,
                            'lift_models': models,
                            'source': 'doc_expansion'
                        })
                        seen_fact_ids.add(fact_id)
                        expansion_count += 1

                if expansion_count:
                    logger.info(f"Document fact expansion added {expansion_count} facts from {doc_filenames}")
            finally:
                cur.close()
                conn.close()
    except Exception as e:
        logger.warning(f"Document fact expansion failed: {e}")

    # Clean up internal fields before returning
    for r in results:
        # Remove internal id field
        if "id" in r:
            del r["id"]
    # Scope-qualifier penalty — demote chunks from the wrong area of the lift
    if SCOPE_QUALIFIER_AVAILABLE:
        results = apply_scope_penalty(results, semantic_query)

    for r in results:
        # Remove internal scoring fields
        if "rrf_score" in r:
            del r["rrf_score"]
        if "similarity" in r:
            del r["similarity"]
        if "source" in r:
            del r["source"]

    logger.info(f"Hybrid search returned {len(results[:limit])} results")
    return results[:limit]


def search_documents(query: str, lift_model: Optional[str] = None, limit: int = 5) -> list[dict]:
    """Search documents using PostgreSQL full-text search with multi-query strategy.

    NOTE: This is the legacy search function. For new code, prefer
    search_documents_hybrid() with interceptor results for better quality.

    Implements multi-query search:
    1. Primary search with the original/expanded query
    2. Fallback searches with query variants (key terms, symptom-focused)

    Then applies three-tier model filtering:
    1. Exact model match
    2. Same model family (if < 3 results)
    3. All models (if still < 3 results)

    Cross-model results are flagged for display.
    """
    conn = get_db_connection()
    cur = conn.cursor()

    results = []
    seen_ids = set()

    # Generate search variants for multi-query strategy
    query_variants = _generate_search_variants(query)
    all_queries = [query] + query_variants  # Primary query first, then variants

    logger.info(f"Multi-query search with {len(all_queries)} variants: {all_queries[:3]}")

    # Tier 1: Exact model match (if model specified)
    if lift_model:
        # Try each query variant until we have enough results
        for q in all_queries:
            if len(results) >= limit:
                break

            remaining = limit - len(results)
            new_results = _run_single_search(cur, q, [lift_model], remaining, seen_ids)
            results.extend(new_results)

            if new_results:
                logger.info(f"Query '{q[:50]}...' found {len(new_results)} results")

        # Tier 2: Family search (if not enough results)
        if len(results) < MIN_RESULTS_THRESHOLD:
            family = get_model_family(lift_model)
            if family:
                other_family = [m for m in family if m != lift_model]
                if other_family:
                    # Try query variants on family
                    for q in all_queries:
                        if len(results) >= limit:
                            break

                        remaining = limit - len(results)
                        family_results = _run_single_search(cur, q, other_family, remaining, seen_ids)
                        for r in family_results:
                            r["cross_model"] = True
                        results.extend(family_results)

        # Tier 3: All models (if still not enough results)
        if len(results) < MIN_RESULTS_THRESHOLD:
            for q in all_queries:
                if len(results) >= limit:
                    break

                remaining = limit - len(results)
                all_results = _run_single_search(cur, q, None, remaining, seen_ids)
                for r in all_results:
                    r["cross_model"] = True
                results.extend(all_results)
    else:
        # No model specified - search all with multi-query
        for q in all_queries:
            if len(results) >= limit:
                break

            remaining = limit - len(results)
            new_results = _run_single_search(cur, q, None, remaining, seen_ids)
            results.extend(new_results)

    # Fallback: If no results and model specified, return recent facts for that model
    if len(results) == 0 and lift_model:
        logger.info(f"No text search results, falling back to model facts for {lift_model}")
        cur.execute("""
            SELECT f.id, f.lift_models, d.filename, d.file_type, f.page, f.content
            FROM facts f
            LEFT JOIN documents d ON f.document_id = d.id
            WHERE f.lift_models && %s
            ORDER BY f.id DESC
            LIMIT %s
        """, ([lift_model], limit))

        for row in cur.fetchall():
            lift_models = row[1] if row[1] else []
            results.append({
                "id": row[0],
                "lift_model": lift_models[0] if lift_models else None,
                "filename": row[2],
                "file_type": row[3],
                "page_number": row[4] or 1,
                "content": row[5],
                "rank": 0,
                "cross_model": False,
                "fallback": True
            })

    cur.close()
    conn.close()

    # Re-rank results for quality
    results = _rerank_results(results, query)

    # Remove internal id field before returning
    for r in results:
        if "id" in r:
            del r["id"]

    logger.info(f"Found {len(results)} document results for query: {query}")
    return results


def search_images(query: str, lift_model: Optional[str] = None, limit: int = 3) -> list[dict]:
    """Search images by description."""
    conn = get_db_connection()
    cur = conn.cursor()

    if lift_model:
        cur.execute("""
            SELECT lift_models, filename, relative_path, description, width, height
            FROM images
            WHERE description_tsv @@ plainto_tsquery('english', %s)
              AND lift_models && %s
            ORDER BY ts_rank_cd(description_tsv, plainto_tsquery('english', %s), 32) DESC
            LIMIT %s
        """, (query, [lift_model], query, limit))
    else:
        cur.execute("""
            SELECT lift_models, filename, relative_path, description, width, height
            FROM images
            WHERE description_tsv @@ plainto_tsquery('english', %s)
            ORDER BY ts_rank_cd(description_tsv, plainto_tsquery('english', %s), 32) DESC
            LIMIT %s
        """, (query, query, limit))

    results = []
    for row in cur.fetchall():
        lift_models = row[0] if row[0] else []
        results.append({
            "lift_model": lift_models[0] if lift_models else None,
            "filename": row[1],
            "relative_path": row[2],
            "description": row[3],
            "width": row[4],
            "height": row[5],
            "full_path": os.path.join(settings.MANUALS_DIR, row[2])
        })

    cur.close()
    conn.close()

    logger.info(f"Found {len(results)} image results for query: {query}")
    return results


# Troubleshooting knowledge for fallback context
TROUBLESHOOTING_CATEGORIES = {
    'door': {
        'common_issues': [
            'Door interlock not engaging - check alignment, mounting mode settings',
            'Door not closing fully - check obstruction sensors, door operator timing',
            'Lock not releasing - check power supply to lock, wiring continuity',
        ],
        'error_codes': ['W01', 'W02', 'W03', 'E01', 'E02'],
        'check_points': ['Door operator power supply', 'Interlock alignment', 'Mounting mode parameter', 'Safety circuit continuity'],
    },
    'motor': {
        'common_issues': [
            'Motor not starting - check contactor, thermal overload, phase sequence',
            'Motor running slow - check inverter parameters, encoder feedback',
            'Motor overheating - check ventilation, load, duty cycle',
        ],
        'error_codes': ['E05', 'E06', 'E07', 'E10', 'E11'],
        'check_points': ['Main contactor', 'Thermal overload', 'Phase voltage', 'Encoder connection', 'Inverter parameters'],
    },
    'controller': {
        'common_issues': [
            'Controller not responding - check power supply, main fuse',
            'Display blank - check display cable, power to controller',
            'Errors on startup - check parameter settings, safety circuit',
        ],
        'error_codes': ['E00', 'E20', 'E21', 'E22', 'E23'],
        'check_points': ['24V power supply', 'Main fuse', 'Safety circuit', 'Parameter backup', 'Firmware version'],
    },
    'leveling': {
        'common_issues': [
            'Poor leveling accuracy - check encoder, slow-down switches',
            'Overrunning floors - check brake, deceleration parameters',
            'Hunting at floor - check leveling sensor, speed parameters',
        ],
        'error_codes': ['E12', 'E13', 'E14', 'E15'],
        'check_points': ['Floor sensor alignment', 'Encoder connection', 'Brake adjustment', 'Deceleration distance'],
    },
    'safety': {
        'common_issues': [
            'Safety circuit open - check all safety devices in sequence',
            'Emergency stop active - check e-stop buttons, key switches',
            'Overload fault - check load weighing device, threshold setting',
        ],
        'error_codes': ['E30', 'E31', 'E32', 'E33'],
        'check_points': ['Safety chain continuity', 'E-stop buttons', 'Limit switches', 'Overload sensor'],
    },
}


def _get_fallback_troubleshooting_context(query: str, lift_model: str = None) -> str:
    """
    Generate fallback troubleshooting context when no direct matches found.

    Provides Claude with common troubleshooting categories and checkpoints
    based on keywords detected in the query.
    """
    query_lower = query.lower()
    relevant_categories = []

    # Detect relevant categories from query
    category_keywords = {
        'door': ['door', 'lock', 'latch', 'interlock', 'gate', 'entrance', 'closing', 'opening'],
        'motor': ['motor', 'pump', 'drive', 'moving', 'travel', 'speed', 'slow', 'not moving'],
        'controller': ['controller', 'board', 'display', 'screen', 'menu', 'parameter', 'setting'],
        'leveling': ['level', 'floor', 'stop', 'position', 'accuracy', 'overrun'],
        'safety': ['safety', 'emergency', 'e-stop', 'overload', 'alarm', 'stuck'],
    }

    for category, keywords in category_keywords.items():
        if any(kw in query_lower for kw in keywords):
            relevant_categories.append(category)

    if not relevant_categories:
        # Default to most common categories
        relevant_categories = ['door', 'controller', 'safety']

    # Build context
    context_parts = []
    context_parts.append(
        f"[TROUBLESHOOTING GUIDE - No exact match found for query. "
        f"Using general troubleshooting knowledge for {lift_model or 'this lift type'}.]"
    )

    for cat in relevant_categories[:2]:  # Limit to 2 categories
        info = TROUBLESHOOTING_CATEGORIES.get(cat, {})
        if info:
            context_parts.append(f"\n**{cat.upper()} TROUBLESHOOTING:**")

            if info.get('common_issues'):
                context_parts.append("Common issues:")
                for issue in info['common_issues'][:3]:
                    context_parts.append(f"  - {issue}")

            if info.get('error_codes'):
                context_parts.append(f"Related error codes: {', '.join(info['error_codes'])}")

            if info.get('check_points'):
                context_parts.append("Key checkpoints: " + ', '.join(info['check_points'][:4]))

    context_parts.append(
        "\n[Use this information to guide troubleshooting. "
        "Ask user to check controller display for specific error codes.]"
    )

    return '\n'.join(context_parts)


def get_fallback_results(query: str, lift_model: str = None) -> List[dict]:
    """
    Generate fallback results when search returns empty.

    Returns results in the same dict format as normal search results,
    ensuring consistent handling downstream.

    Args:
        query: Original search query
        lift_model: Optional lift model context

    Returns:
        List with single fallback result dict
    """
    fallback_text = _get_fallback_troubleshooting_context(query, lift_model)
    return [{
        "filename": "troubleshooting_guide",
        "page_number": 0,
        "lift_model": lift_model or "General",
        "content": fallback_text,
        "chunk_type": "fallback",
        "cross_model": False,
        "fallback": True
    }]


def format_context(doc_results: list[dict], image_results: list[dict] = None,
                   query: str = None, lift_model: str = None) -> str:
    """Format search results into context for Claude with clear Source IDs.

    Each chunk is formatted as:
    [Source ID: {filename} | Page: {page_number} | Model: {lift_model}]
    {content}

    This enables precise citations like [Source: E3_Manual.pdf].
    Cross-model results are flagged for verification.
    """
    context_parts = []

    # Check if all results are fallback (no text match found)
    has_fallback = any(r.get("fallback", False) for r in doc_results) if doc_results else False

    # If no results at all, generate fallback results in consistent format
    if not doc_results and query:
        doc_results = get_fallback_results(query, lift_model)

    elif has_fallback and doc_results:
        context_parts.append(
            "[SYSTEM NOTE: No exact documentation match. Showing general model information.]"
        )

    # Document context with clear Source IDs
    if doc_results:
        for r in doc_results:
            is_cross = r.get("cross_model", False)
            is_fallback = r.get("fallback", False)
            model = r.get('lift_model', 'Unknown')

            # Build source ID header with optional chunk type indicator
            chunk_type_indicator = ""
            if r.get("chunk_type") == "table":
                chunk_type_indicator = " [TABLE DATA]"

            source_header = f"[Source ID: {r['filename']} | Page: {r['page_number']} | Model: {model}{chunk_type_indicator}]"

            # Add cross-model warning if applicable
            # E3 has different architecture than standard Elfo - add special warning
            if is_cross and lift_model == "E3" and model in ["Elfo", "Elfo 2"]:
                source_header += "\n[CROSS-MODEL WARNING: E3 has different architecture than standard Elfo. Verify procedures before applying.]"
            elif is_cross:
                source_header += f"\n[CROSS-MODEL: This info is from {model} - verify for {lift_model}]"
            elif is_fallback:
                source_header += "\n[GENERAL INFO: Not a direct match for your query.]"

            context_parts.append(f"{source_header}\n{r['content']}")

    # Image references with Source IDs
    if image_results:
        context_parts.append("[RELEVANT IMAGES]")
        for r in image_results:
            context_parts.append(
                f"[Source ID: {r['filename']} | Model: {r['lift_model']}]\n"
                f"Image: {r['description']} ({r['width']}x{r['height']})"
            )

    return "\n\n---\n\n".join(context_parts) if context_parts else ""


def get_all_extensions():
    """Get list of all supported file extensions."""
    all_ext = []
    for extensions in SUPPORTED_EXTENSIONS.values():
        all_ext.extend(extensions)
    return all_ext


def ingest_all_manuals(progress_callback=None):
    """
    Ingest all files from the manuals directory.

    Args:
        progress_callback: Optional function(current, total, filename) for progress updates
    """
    manuals_dir = settings.MANUALS_DIR

    if not os.path.exists(manuals_dir):
        logger.warning(f"Manuals directory not found: {manuals_dir}")
        return {"error": "Manuals directory not found"}

    # Collect all files first
    all_files = []
    for model_dir in os.listdir(manuals_dir):
        model_path = os.path.join(manuals_dir, model_dir)
        if not os.path.isdir(model_path):
            continue

        for filename in os.listdir(model_path):
            filepath = os.path.join(model_path, filename)
            if os.path.isfile(filepath) and get_file_type(filepath):
                all_files.append((filepath, model_dir, filename))

    total = len(all_files)
    success_count = 0
    error_count = 0

    for i, (filepath, model_dir, filename) in enumerate(all_files):
        file_type = get_file_type(filepath)
        relative_path = os.path.relpath(filepath, manuals_dir)

        if progress_callback:
            progress_callback(i + 1, total, filename)

        try:
            chunks = ingest_file(filepath, model_dir)
            update_index_status(model_dir, filename, file_type, relative_path, 'success', chunks)
            success_count += 1
            logger.info(f"[{i+1}/{total}] Indexed: {filename} ({chunks} chunks)")
        except Exception as e:
            error_msg = str(e)
            update_index_status(model_dir, filename, file_type, relative_path, 'error', 0, error_msg)
            error_count += 1
            logger.error(f"[{i+1}/{total}] Failed: {filename} - {error_msg}")

    return {
        "total": total,
        "success": success_count,
        "errors": error_count
    }
